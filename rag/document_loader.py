"""
文档加载器
支持 PDF、Word(.docx)、Markdown(.md) 三种格式
输出统一的 Document 对象列表
"""

import os
import sys
if sys.stdout.encoding != 'utf-8':
    sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
from typing import List
from langchain_core.documents import Document

class DocumentLoader:
    """统一文档加载器，根据文件扩展名自动选择解析器"""
    #支持的文件类型
    SUPPORTED_EXTENSIONS={
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.md': 'markdown',
        '.markdown': 'markdown'
    }

    def __init__(self,knowledge_base_dir:str):
        """
        Args:
            knowledge_base_dir: 知识库根目录路径
        """
        self.knowledge_base_dir=knowledge_base_dir

    def load_all(self)->List[Document]:
        """加载所有文档，统一转为 Markdown 格式的 Document"""

        """ os.walk() 生成器每次 yield 一个三元组 (dirpath, dirnames, filenames)，遍历整个目录树。

root — 当前遍历到的目录路径（字符串）。第一趟是根目录，后续会递归进入子目录。
_（即 dirnames） — 当前目录下所有的子目录名列表（字符串列表）。用 _ 表示你不需要用到它。
files — 当前目录下所有的文件名列表（字符串列表），不含路径。
举个例子，假设目录结构是：


knowledge_base/
├── intro.txt
├── docs/
│   └── guide.pdf
└── images/
    └── logo.png
os.walk("knowledge_base") 会依次 yield：


("knowledge_base",          ["docs", "images"],  ["intro.txt"])
("knowledge_base/docs",     [""],                ["guide.pdf"])
("knowledge_base/images",   [""],                ["logo.png"]) """
        all_documents=[]

        if not os.path.exists(self.knowledge_base_dir):
            print(f"⚠️ 知识库目录不存在：{self.knowledge_base_dir}")
            return all_documents
        
        for root,_,files in os.walk(self.knowledge_base_dir):
            for filename in files:
                file_path=os.path.join(root,filename)
                """
                os.path.splitext(filename) 把文件名拆成 (名称, 扩展名) 两部分：
os.path.splitext("guide.pdf")       # 返回 ("guide", ".pdf")
os.path.splitext("report.TXT")      # 返回 ("report", ".TXT")
os.path.splitext("noextension")     # 返回 ("noextension", "")"""
                ext=os.path.splitext(filename)[1].lower()

                if ext not in self.SUPPORTED_EXTENSIONS:
                    continue

                try:
                    documents=self.load_single_file(file_path,ext)
                    all_documents.extend(documents)
                    print(f"✅ 已加载：{filename}（{len(documents)} 个片段）")
                except Exception as e:
                    print(f"❌ 加载文件 {filename} 失败：{e}")
        
        print(f"\n📚 总计加载 {len(all_documents)} 个文档片段")
        return all_documents
    
    def load_single_file(self,file_path:str,ext:str)->List[Document]:
        """根据扩展名调用对应加载器，统一输出 Markdown"""
        if ext=='.pdf':
            markdown_content=self. _pdf_to_markdown(file_path)
        elif ext=='.docx':
            markdown_content = self._word_to_markdown(file_path)
        elif ext in ('.md', '.markdown'):
            markdown_content = self._read_markdown(file_path)
        else:
            raise ValueError(f"不支持的文件格式：{ext}")
        
        if not markdown_content.strip():
            return []
        
        # 整篇文档作为一个 Document，后续由分块器按 Markdown 标题切分
        return [Document(
            page_content=markdown_content,
            metadata={
                "source": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": ext.lstrip('.'),
            }
        )]
    
    # ─── PDF 转 Markdown ───
    def _pdf_to_markdown(self,file_path:str)->str:
        """PDF → Markdown，根据字体大小推断标题层级"""
        import fitz
        doc=fitz.open(file_path)
        markdown_parts=[]

        """range(len(doc)) — 遍历每一页（len(doc) 返回总页数）
        page=doc[page_num] — 获取指定页的对象
        page.get_text("dict")["blocks"] — 以字典格式提取页面文本块，blocks 是一个列表，
        每个元素包含该块的类型（文字/图片/表格等）和文本内容
        """

        for page_num in range(len(doc)):
            page=doc[page_num]
            blocks=page.get_text("dict")["blocks"]
            page_markdown=self._blocks_to_markdown(blocks)

            if page_markdown.strip():
                markdown_parts.append(page_markdown)
        
        doc.close()
        return "\n\n".join(markdown_parts)
    

    def _blocks_to_markdown(self,blocks:List)->str:
        """将 PDF 文本块转换为 Markdown，根据字体大小推断标题"""
        lines=[]
        for block in blocks:
            if block["type"] != 0:  # 非文字块，跳过
                continue

            for line in block["lines"]:
                text="".join([span["text"]for span in line["spans"]])

                if not text.strip():  # 空行，跳过
                    lines.append("")
                    continue

                """
                block["type"] != 0

fitz 中 type=0 表示文本块，其他（图片、表格等）跳过。
"".join([span["text"] for span in line["spans"]])

一行可能由多个 <span> 组成（比如部分加粗、不同字体），这里把它们拼接成完整文本。
spans[0].get("size", 12)

取该行第一个 span 的字体大小（单位是 pt），默认 12。
spans[0].get("flags", 0) & 2**3

这是判断是否加粗的位运算：
flags 是一个整数字段，每一位代表一个属性
2**3 = 8，第 4 位（从 0 开始）表示 bold
& 8 结果非零 → 是粗体
字体大小 → 标题层级的规则：

条件	Markdown
size ≥ 20	# 一级标题
size ≥ 16	## 二级标题
size ≥ 14 且加粗	### 三级标题
其他	普通正文
核心思路就是： PDF 里没有 Markdown 语法，但通常标题的字号更大。
代码通过字体大小（加粗辅助判断）来"猜测"哪些行应该是 # 标题。
                """

                # 根据字体大小推断标题层级
                spans=line["spans"]
                if spans:
                    font_size=spans[0].get("size", 12)
                    is_bold=bool(spans[0].get("flags", 0) & 2**3)

                    if font_size >= 20:
                        text=f"# {text}"
                    elif font_size >= 16:
                        text=f"## {text}"
                    elif font_size >= 14 and is_bold:
                        text=f"### {text}"

                lines.append(text)
            return "\n".join(lines)
        
        

    # ─── Word 转 Markdown ───
    def _word_to_markdown(self,file_path:str)->str:
        """Word(.docx) → Markdown，保留标题、列表、表格"""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        markdown_parts=[]
        """"
        DocxDocument(file_path) — 用 python-docx 打开 .docx 文件
        doc.paragraphs — 遍历所有段落（文字、标

        para.style.name.lower() — 获取段落样式名称，如 "Heading 1"、"List Paragraph"
样式匹配逻辑：
样式名包含	Markdown
heading 1 / 标题 1	#
heading 2 / 标题 2	##
heading 3 / 标题 3	###
其他 heading	####
list	- （无序列表）
其他	纯文本
        """

        for para in doc.paragraphs:
            text=para.text.strip()
            if not text:
                markdown_parts.append("")
                continue

            # 根据 Word 样式推断 Markdown 格式
            style_name=para.style.name.lower() if para.style else ""

            if style_name.startswith("heading 1") or style_name.startswith("标题 1"):
                markdown_parts.append(f"# {text}")
            elif style_name.startswith("heading 2") or style_name.startswith("标题 2"):
                markdown_parts.append(f"## {text}")
            elif style_name.startswith("heading 3") or style_name.startswith("标题 3"):
                markdown_parts.append(f"### {text}")
            elif style_name.startswith("heading"):
                markdown_parts.append(f"#### {text}")
            elif para.style and "list" in style_name:
                markdown_parts.append(f"- {text}")
            else:
                markdown_parts.append(text)

        #处理表格
        for table in doc.tables:
            markdown_parts.append("")
            markdown_parts.append(self._table_to_markdown(table))
        
        return "\n".join(markdown_parts)
    
    def _table_to_markdown(self,table)->str:
        """将 Word 表格转换为 Markdown 表格"""
        rows=[]
        for row in table.rows:
            cells=[cell.text.strip().replace("\n","") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        
        if not rows:
            return ""
        # 插入表头分隔线
        col_count=len(table.rows[0].cells) if table.rows else 1
        separator="| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1,separator)

        return "\n".join(rows)
    

    # ─── Markdown 原样读取 ───
    def _read_markdown(self,file_path:str)->str:
        """Markdown 文件直接读取，原样保留"""
        with open(file_path,"r",encoding="utf-8") as f:
            return f.read()








                
